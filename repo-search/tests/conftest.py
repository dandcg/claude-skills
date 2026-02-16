"""Shared fixtures for repo-search tests.

Provides programmatically generated test documents in multiple formats
(Markdown, PDF, DOCX, XLSX) and a pre-ingested vector database fixture.
"""

from pathlib import Path

import pytest


# ---------------------------------------------------------------------------
# Markdown content helpers
# ---------------------------------------------------------------------------

_Q4_REVENUE_MD = """\
---
**Date:** 2025-01-15
**Status:** Final
---

# Q4 Revenue Report

## Summary

The fourth quarter of 2025 showed strong performance across all business
segments. Total revenue reached $2.3M, exceeding our forecast by 12%.
This growth was primarily driven by expansion in the APAC region and
increased adoption of our enterprise product tier. Customer retention
remained high at 94%, reflecting the value our platform delivers.

## Regional Breakdown

### North America

North American revenue contributed $1.1M to the quarterly total,
representing a 8% increase over Q3. Enterprise deals in financial
services and healthcare drove this growth, with average deal sizes
increasing by 15%. The sales team closed 23 new enterprise accounts
during the quarter, bringing total enterprise customers to 187.

### APAC

The APAC region delivered $0.7M in revenue, marking a 25% quarter-
over-quarter increase. Japan and Australia were the strongest markets,
with combined revenue of $0.5M. Our localization efforts and regional
partnership with TechDistribute Co. significantly accelerated adoption.

### EMEA

EMEA contributed $0.5M, a steady 5% growth from Q3. The UK and Germany
remained our largest markets in the region. Regulatory compliance
features released in October helped win several government contracts.

## Outlook

Looking ahead to Q1 2026, we project revenue of $2.5M based on current
pipeline. Key initiatives include expanding the partner channel in APAC,
launching the self-serve enterprise tier, and entering the Latin American
market. Headcount will increase by 12 across sales and engineering to
support these growth plans.
"""

_EXERCISE_ROUTINE_MD = """\
---
**Added:** 2025-03-10
---

# Exercise Routine

## Morning Workout

The morning workout begins at 6:30 AM with a 10-minute dynamic warm-up
consisting of jumping jacks, high knees, and arm circles. This is
followed by a 30-minute strength training circuit that alternates
between upper body and lower body exercises. On Monday, Wednesday, and
Friday, the focus is on compound movements like squats, deadlifts, and
bench press. Tuesday and Thursday emphasize isolation exercises and
core work including planks, Russian twists, and leg raises.

## Evening Stretching

The evening stretching routine takes 20 minutes and focuses on
flexibility and recovery. It includes hamstring stretches, hip
flexor stretches, shoulder mobility work, and a 5-minute guided
meditation for relaxation. Foam rolling is incorporated on days
following heavy lifting sessions to aid muscle recovery.
"""


def _generate_python_best_practices_md() -> str:
    """Generate a long markdown document with 20 sections, each ~20 sentences."""
    sections = [
        ("Code Style and Formatting",
         "Consistent code style makes projects easier to maintain. "
         "Use PEP 8 as the baseline for all Python projects. "
         "Configure your editor to auto-format on save using tools like Black. "
         "Black enforces a deterministic style that eliminates debates about formatting. "
         "Line length should be limited to 88 characters when using Black. "
         "Use isort to keep import statements organized and grouped correctly. "
         "Separate standard library, third-party, and local imports with blank lines. "
         "Avoid wildcard imports as they pollute the namespace. "
         "Use absolute imports for clarity in larger projects. "
         "Relative imports are acceptable within packages for internal modules. "
         "Consistent naming conventions improve readability significantly. "
         "Use snake_case for functions and variables throughout the codebase. "
         "Class names should follow PascalCase conventions. "
         "Constants should be UPPER_SNAKE_CASE at the module level. "
         "Private attributes and methods should start with an underscore. "
         "Avoid single-character variable names except in comprehensions. "
         "Type hints serve as documentation and enable static analysis. "
         "Configure pre-commit hooks to enforce style automatically. "
         "Flake8 catches additional issues that Black does not address. "
         "Pylint provides deeper analysis but requires more configuration."),

        ("Virtual Environments",
         "Always use virtual environments to isolate project dependencies. "
         "The venv module is built into Python 3 and sufficient for most needs. "
         "Create a virtual environment with python -m venv .venv in each project. "
         "Activate the environment before installing any packages. "
         "Never install packages into the system Python installation. "
         "Use pip freeze to capture exact dependency versions. "
         "Pin all dependency versions in requirements.txt for reproducibility. "
         "Consider using pip-tools for better dependency resolution. "
         "The pip-compile command generates pinned requirements from abstract specs. "
         "Poetry and PDM offer modern dependency management alternatives. "
         "They handle virtual environments and lock files automatically. "
         "Conda is useful for scientific computing with non-Python dependencies. "
         "Docker containers provide even stronger isolation for deployment. "
         "Document the Python version requirement in pyproject.toml. "
         "Use pyenv to manage multiple Python versions on development machines. "
         "Test against the minimum supported Python version in CI. "
         "Keep development and production dependencies separate. "
         "Requirements-dev.txt should include testing and linting tools. "
         "Regularly update dependencies to get security patches. "
         "Use dependabot or renovate to automate dependency update PRs."),

        ("Error Handling",
         "Proper error handling prevents silent failures in production. "
         "Catch specific exceptions rather than using bare except clauses. "
         "The Exception base class is too broad for most catch blocks. "
         "Create custom exception classes for domain-specific errors. "
         "Custom exceptions should inherit from appropriate built-in exceptions. "
         "Use exception chaining with raise from to preserve context. "
         "Log exceptions with full tracebacks for debugging purposes. "
         "The logging module provides structured exception logging. "
         "Avoid using exceptions for control flow in normal operation. "
         "EAFP style is Pythonic but should not replace proper validation. "
         "Validate inputs at system boundaries to fail fast. "
         "Use assertions for invariants that should never be violated. "
         "Assertions are removed in optimized mode so do not use them for validation. "
         "Context managers with __exit__ can handle cleanup on exceptions. "
         "The contextlib module simplifies creating context managers. "
         "Retry logic should use exponential backoff for transient failures. "
         "The tenacity library provides robust retry decorators. "
         "Always include meaningful error messages in exception strings. "
         "Error messages should describe what went wrong and suggest fixes. "
         "Structured error responses in APIs help clients handle failures."),

        ("Testing Fundamentals",
         "Testing is essential for maintaining code quality over time. "
         "Write tests before fixing bugs to prevent regressions. "
         "Unit tests should be fast, isolated, and deterministic. "
         "Each test should verify a single behavior or scenario. "
         "Use descriptive test names that explain the expected behavior. "
         "The Arrange-Act-Assert pattern structures tests clearly. "
         "Fixtures reduce duplication in test setup code. "
         "Pytest fixtures are more flexible than unittest setUp methods. "
         "Parametrize tests to cover multiple input combinations. "
         "Mocking external dependencies keeps unit tests fast. "
         "The unittest.mock module provides comprehensive mocking support. "
         "Patch at the point of use rather than the point of definition. "
         "Integration tests verify components work together correctly. "
         "Keep integration tests separate from unit tests for speed. "
         "Use markers to categorize tests by type and speed. "
         "Aim for high coverage but do not chase 100% blindly. "
         "Coverage tools identify untested code paths. "
         "Mutation testing reveals weaknesses in your test suite. "
         "Property-based testing with Hypothesis finds edge cases. "
         "Test data builders make creating test objects convenient."),

        ("Documentation",
         "Good documentation saves time for every team member. "
         "Docstrings should describe what a function does and why. "
         "Use Google or NumPy style docstrings consistently. "
         "Include parameter types and return types in docstrings. "
         "Document exceptions that a function may raise. "
         "Module-level docstrings explain the purpose of the module. "
         "Keep README files current with setup and usage instructions. "
         "Architecture decision records capture important design choices. "
         "API documentation should include examples for every endpoint. "
         "Use Sphinx or MkDocs to generate documentation from docstrings. "
         "Type hints serve as a form of executable documentation. "
         "Inline comments should explain why, not what the code does. "
         "Remove commented-out code rather than leaving it as documentation. "
         "Version your documentation alongside the code it describes. "
         "Changelog entries help users understand what changed between versions. "
         "Include migration guides for breaking changes. "
         "Tutorial-style documentation helps new users get started quickly. "
         "Reference documentation should be comprehensive and searchable. "
         "Diagrams in documentation clarify complex system interactions. "
         "Review documentation as part of the code review process."),

        ("Project Structure",
         "A well-organized project structure makes navigation intuitive. "
         "Use the src layout to prevent accidental imports of uninstalled code. "
         "The src directory contains only the package source code. "
         "Tests live in a top-level tests directory parallel to src. "
         "Configuration files belong in the project root. "
         "Use pyproject.toml as the single source of project metadata. "
         "The setup.cfg format is being replaced by pyproject.toml. "
         "Keep scripts and tools in a dedicated scripts directory. "
         "Data files should be in a separate data or resources directory. "
         "Use importlib.resources to access package data files. "
         "Avoid deeply nested directory structures that slow navigation. "
         "Each package should have a clear single responsibility. "
         "Circular imports indicate poor module boundaries. "
         "Use __init__.py to define the public API of each package. "
         "Keep __init__.py files minimal to reduce import side effects. "
         "Namespace packages allow splitting code across multiple directories. "
         "Monorepos can house multiple related packages effectively. "
         "Use a consistent naming convention for modules across the project. "
         "Group related functionality into cohesive packages. "
         "The project structure should mirror the logical architecture."),

        ("Dependency Injection",
         "Dependency injection makes code more testable and flexible. "
         "Pass dependencies as parameters instead of creating them internally. "
         "Constructor injection is the most common form in Python. "
         "Function parameter injection works well for simpler cases. "
         "Avoid global state that creates hidden dependencies between modules. "
         "The dependency inversion principle guides good interface design. "
         "High-level modules should depend on abstractions not concretions. "
         "Protocol classes in Python define structural interfaces. "
         "Abstract base classes enforce interface contracts at runtime. "
         "Use factories to create complex objects with many dependencies. "
         "The injector library provides a formal DI container for Python. "
         "Simpler projects can use manual dependency wiring effectively. "
         "Configuration should be injected rather than read from globals. "
         "Environment variables are a common source of configuration. "
         "The settings pattern centralizes configuration management. "
         "Pydantic Settings validates configuration from environment variables. "
         "Separate construction of objects from their usage. "
         "The composition root is where all dependencies get wired together. "
         "Test doubles replace real dependencies in unit tests. "
         "Dependency injection reduces coupling between system components."),

        ("Async Programming",
         "Asyncio enables concurrent IO operations in a single thread. "
         "Use async def and await for coroutine-based concurrency. "
         "The event loop manages scheduling of coroutines efficiently. "
         "Asyncio is ideal for IO-bound workloads like web servers. "
         "CPU-bound work should use multiprocessing not asyncio. "
         "Aiohttp provides async HTTP client and server functionality. "
         "FastAPI leverages async for high-performance web APIs. "
         "Database access can be async with libraries like asyncpg. "
         "Use asyncio.gather to run multiple coroutines concurrently. "
         "Task groups in Python 3.11 provide structured concurrency. "
         "Structured concurrency prevents orphaned tasks and lost exceptions. "
         "Cancel scopes allow clean timeout handling for async operations. "
         "Async context managers handle resource lifecycle in async code. "
         "Async iterators enable streaming data processing patterns. "
         "Avoid mixing sync and async code without careful bridging. "
         "Use asyncio.to_thread to run blocking code from async contexts. "
         "Semaphores limit concurrency to prevent resource exhaustion. "
         "Connection pools manage shared resources in async applications. "
         "Testing async code requires pytest-asyncio or similar helpers. "
         "Profile async code with asyncio debug mode to find slowdowns."),

        ("Security Practices",
         "Security should be considered throughout the development process. "
         "Never store secrets in source code or version control. "
         "Use environment variables or secret management services for credentials. "
         "The secrets module generates cryptographically strong random values. "
         "Validate and sanitize all user input before processing. "
         "SQL injection is prevented by using parameterized queries. "
         "ORMs like SQLAlchemy handle query parameterization automatically. "
         "Cross-site scripting is prevented by escaping output properly. "
         "Use HTTPS for all network communication in production. "
         "Pin SSL certificate authorities for sensitive API connections. "
         "Hash passwords with bcrypt or argon2 never with MD5 or SHA. "
         "Implement rate limiting to prevent brute force attacks. "
         "Use CSRF tokens to protect form submissions. "
         "Set appropriate CORS headers for web APIs. "
         "Regularly update dependencies to patch known vulnerabilities. "
         "Run Bandit to detect common security issues in Python code. "
         "Safety checks installed packages against known vulnerability databases. "
         "Use least privilege principles for service accounts. "
         "Audit logging tracks security-relevant events for investigation. "
         "Code review should include security considerations explicitly."),

        ("Performance Optimization",
         "Measure performance before optimizing to avoid premature optimization. "
         "Use cProfile to identify the actual bottlenecks in your code. "
         "Line profiler reveals which lines consume the most time. "
         "Memory profiler helps identify memory leaks and excessive allocation. "
         "Algorithmic improvements provide the largest performance gains. "
         "Choose appropriate data structures for your access patterns. "
         "Sets provide O(1) membership testing compared to O(n) for lists. "
         "Dictionaries offer fast key-value lookups. "
         "Use generators for large datasets to reduce memory consumption. "
         "List comprehensions are faster than equivalent for loops. "
         "The functools.lru_cache decorator caches expensive function results. "
         "NumPy vectorizes operations for significant speedups on arrays. "
         "Pandas optimizes data manipulation for tabular data. "
         "Batch database operations rather than executing them one by one. "
         "Connection pooling reduces overhead for database connections. "
         "Use asyncio for IO-bound concurrency improvements. "
         "Multiprocessing handles CPU-bound parallel workloads. "
         "Cython compiles Python to C for performance-critical sections. "
         "PyPy provides significant speedups for pure Python code. "
         "Load testing verifies performance under realistic conditions."),

        ("Logging Best Practices",
         "Logging provides visibility into application behavior in production. "
         "Use the standard logging module for consistent log management. "
         "Configure logging at application startup not in library code. "
         "Libraries should use logging.getLogger(__name__) for their loggers. "
         "Use appropriate log levels: DEBUG INFO WARNING ERROR CRITICAL. "
         "DEBUG logs are for detailed diagnostic information. "
         "INFO logs record normal operational events. "
         "WARNING logs indicate potential issues that need attention. "
         "ERROR logs capture failures that affect functionality. "
         "CRITICAL logs record severe failures that may crash the application. "
         "Structured logging with JSON format enables easier log analysis. "
         "Include correlation IDs in logs to trace requests across services. "
         "Avoid logging sensitive data like passwords or personal information. "
         "Log rotation prevents log files from consuming all disk space. "
         "Centralized log aggregation simplifies debugging distributed systems. "
         "ELK stack and Grafana Loki are popular log aggregation solutions. "
         "Use log sampling in high-throughput systems to reduce volume. "
         "Contextual logging adds request-specific data to all log entries. "
         "Python structlog library provides excellent structured logging. "
         "Monitor log error rates to detect issues proactively."),

        ("Database Access Patterns",
         "Choose the right database abstraction level for your project. "
         "Raw SQL gives maximum control but requires careful management. "
         "SQLAlchemy Core provides a SQL expression language with Python. "
         "SQLAlchemy ORM maps Python classes to database tables. "
         "The Unit of Work pattern manages transaction boundaries. "
         "Repository pattern abstracts data access behind a clean interface. "
         "Use migrations to manage database schema changes over time. "
         "Alembic integrates with SQLAlchemy for migration management. "
         "Always use transactions for operations that must be atomic. "
         "Connection pooling is essential for production database access. "
         "Prepared statements improve performance for repeated queries. "
         "Indexes speed up queries but slow down write operations. "
         "Analyze query plans to optimize slow database operations. "
         "Use database-specific features when portability is not required. "
         "NoSQL databases suit certain access patterns better than SQL. "
         "Redis provides fast caching and session storage. "
         "MongoDB stores flexible document structures effectively. "
         "Use database fixtures in tests for realistic test data. "
         "Separate read and write models for complex query requirements. "
         "Event sourcing captures all changes as immutable events."),

        ("API Design",
         "Design APIs from the consumer perspective first. "
         "RESTful APIs use HTTP methods to express intent clearly. "
         "Use nouns for resource URLs and verbs through HTTP methods. "
         "Consistent naming conventions make APIs predictable. "
         "Version your API to allow backward-compatible evolution. "
         "URL versioning like /v1/resources is the most common approach. "
         "Return appropriate HTTP status codes for all responses. "
         "Use 201 Created for successful resource creation. "
         "Use 404 Not Found only when the resource does not exist. "
         "Pagination prevents returning excessively large result sets. "
         "Cursor-based pagination is more efficient than offset-based. "
         "HATEOAS links help clients discover related resources. "
         "Rate limiting protects APIs from abuse and overload. "
         "Authentication should use industry standards like OAuth 2.0. "
         "API keys are simpler but less secure than token-based auth. "
         "OpenAPI specifications document APIs in a machine-readable format. "
         "Generate client SDKs from OpenAPI specs for convenience. "
         "Validate request payloads with Pydantic models in FastAPI. "
         "GraphQL provides flexible querying for complex data requirements. "
         "gRPC offers high-performance binary protocol for service communication."),

        ("Configuration Management",
         "Externalize configuration from application code. "
         "Environment variables are the standard for twelve-factor apps. "
         "Use dotenv files for local development configuration. "
         "Never commit dotenv files to version control. "
         "Pydantic Settings validates configuration values at startup. "
         "Fail fast if required configuration is missing. "
         "Provide sensible defaults for optional configuration values. "
         "Use different configuration profiles for development and production. "
         "Feature flags enable gradual rollout of new functionality. "
         "LaunchDarkly and Unleash provide feature flag management. "
         "Centralized configuration services help manage distributed systems. "
         "Consul and etcd store distributed configuration reliably. "
         "Configuration changes should not require application restarts. "
         "Hot reloading configuration reduces deployment downtime. "
         "Encrypt sensitive configuration values at rest. "
         "Audit who changed configuration and when for compliance. "
         "Use configuration schemas to validate structure and types. "
         "Document all configuration options and their effects. "
         "Infrastructure as code manages environment configuration. "
         "Terraform and Ansible automate infrastructure configuration."),

        ("Continuous Integration",
         "Continuous integration catches issues early in the development cycle. "
         "Run the full test suite on every pull request automatically. "
         "GitHub Actions provides CI/CD integrated with GitHub repositories. "
         "GitLab CI/CD is built into GitLab for seamless integration. "
         "Jenkins offers maximum flexibility for complex CI pipelines. "
         "Use matrix builds to test against multiple Python versions. "
         "Cache pip dependencies to speed up CI build times. "
         "Run linting and type checking as CI pipeline stages. "
         "Fail the pipeline on any linting or type errors. "
         "Measure test coverage and track trends over time. "
         "Set minimum coverage thresholds to prevent regression. "
         "Build and test Docker images as part of the CI pipeline. "
         "Security scanning should be integrated into CI pipelines. "
         "Artifact publishing happens after all checks pass. "
         "Use semantic versioning for release tagging. "
         "Changelog generation can be automated from commit messages. "
         "Deploy to staging environments automatically for testing. "
         "Production deployments should require manual approval. "
         "Canary deployments reduce risk for production releases. "
         "Monitor error rates after each deployment for quick rollback."),

        ("Type Hints and Static Analysis",
         "Type hints improve code readability and enable static analysis. "
         "Use type hints for function parameters and return values. "
         "The typing module provides generic types and special forms. "
         "Union types express that a value can be one of several types. "
         "Optional is shorthand for Union with None. "
         "Use TypeVar for generic functions that preserve input types. "
         "Protocol classes define structural subtyping interfaces. "
         "Mypy is the most widely used static type checker for Python. "
         "Pyright offers faster type checking with better IDE integration. "
         "Configure strict mode for maximum type safety in new projects. "
         "Gradual typing allows adding types incrementally to existing code. "
         "Use type stubs for third-party libraries without inline types. "
         "The typeshed project maintains stubs for the standard library. "
         "Runtime type checking with beartype catches errors early. "
         "Pydantic uses type hints for data validation at runtime. "
         "Dataclasses leverage type hints for boilerplate-free classes. "
         "TypedDict adds type safety to dictionary-based data structures. "
         "Literal types restrict values to specific constants. "
         "NewType creates distinct types for documentation purposes. "
         "Type aliases simplify complex type expressions."),

        ("Concurrency Patterns",
         "Choose the right concurrency model for your workload type. "
         "Threading suits IO-bound tasks with shared memory requirements. "
         "The GIL limits Python threads to one CPU core at a time. "
         "Multiprocessing bypasses the GIL for CPU-bound parallelism. "
         "Process pools manage worker processes efficiently. "
         "Concurrent.futures provides a unified API for threads and processes. "
         "ThreadPoolExecutor handles IO-bound concurrent tasks well. "
         "ProcessPoolExecutor distributes CPU-bound work across cores. "
         "Asyncio provides cooperative multitasking for IO operations. "
         "Queue-based patterns decouple producers from consumers. "
         "Threading.Queue is thread-safe for producer-consumer patterns. "
         "Locks protect shared state from concurrent modification. "
         "Deadlocks occur when multiple locks are acquired in different orders. "
         "Use context managers for lock acquisition to ensure release. "
         "Events coordinate thread execution without polling. "
         "Barriers synchronize multiple threads at a specific point. "
         "Semaphores limit the number of concurrent resource accesses. "
         "Actor model isolates state within independent actors. "
         "Message passing avoids shared state concurrency issues. "
         "Celery distributes tasks across worker processes and machines."),

        ("Packaging and Distribution",
         "Modern Python packaging uses pyproject.toml as the standard. "
         "The build module creates distribution packages from source. "
         "Wheel files provide faster installation than source distributions. "
         "Setuptools remains the most common build backend. "
         "Flit and Hatch offer simpler alternatives for pure Python packages. "
         "Use semantic versioning to communicate change impact. "
         "Publish packages to PyPI for public distribution. "
         "Private package indexes host internal company packages. "
         "Use twine to upload packages to PyPI securely. "
         "Test package installation in a clean environment before publishing. "
         "Entry points define console scripts and plugin interfaces. "
         "Include a LICENSE file in every distributed package. "
         "Classifiers on PyPI help users discover your package. "
         "Package metadata should include author and homepage information. "
         "Use MANIFEST.in to control which files are included. "
         "Exclude test files and development tools from distributions. "
         "Namespace packages allow multiple packages to share a prefix. "
         "Conda packages handle non-Python dependencies effectively. "
         "Docker images provide complete runtime environments. "
         "Multi-stage Docker builds reduce final image size significantly."),

        ("Design Patterns in Python",
         "Design patterns provide proven solutions to common problems. "
         "The Singleton pattern ensures only one instance exists. "
         "Python modules are natural singletons due to import caching. "
         "The Factory pattern creates objects without specifying exact classes. "
         "Abstract factories create families of related objects. "
         "The Builder pattern constructs complex objects step by step. "
         "Method chaining provides a fluent interface for builders. "
         "The Observer pattern enables loose coupling between components. "
         "Python signals and events implement the observer pattern. "
         "The Strategy pattern selects algorithms at runtime. "
         "First-class functions in Python simplify the strategy pattern. "
         "The Decorator pattern adds behavior to objects dynamically. "
         "Python decorators wrap functions with additional functionality. "
         "The Adapter pattern makes incompatible interfaces compatible. "
         "The Facade pattern simplifies complex subsystem interfaces. "
         "The Iterator pattern provides sequential access to collections. "
         "Python iterators and generators implement this pattern natively. "
         "The Command pattern encapsulates operations as objects. "
         "The State pattern manages object behavior based on internal state. "
         "Choose patterns based on actual needs not theoretical elegance."),

        ("Debugging Techniques",
         "Effective debugging skills save hours of development time. "
         "Read error messages carefully as they usually indicate the problem. "
         "Tracebacks show the call stack from bottom to top. "
         "The pdb debugger allows stepping through code interactively. "
         "Use breakpoint() to set debugger entry points in Python 3.7 plus. "
         "IDE debuggers provide graphical stepping and variable inspection. "
         "Print debugging is simple but should be removed after use. "
         "Logging is a better alternative to print debugging. "
         "Reproduce bugs consistently before attempting to fix them. "
         "Write a failing test that demonstrates the bug first. "
         "Binary search through code changes finds when bugs were introduced. "
         "Git bisect automates finding the commit that introduced a bug. "
         "Rubber duck debugging explains the problem to an inanimate object. "
         "Taking breaks helps see problems from a fresh perspective. "
         "Remote debugging connects debuggers to running production code. "
         "Post-mortem debugging examines state after an exception occurs. "
         "Memory debuggers like objgraph visualize object references. "
         "Network debugging tools like mitmproxy inspect API traffic. "
         "Profile-guided debugging focuses on actual performance bottlenecks. "
         "Document the root cause and fix for future reference."),
    ]

    header = """\
---
**Date:** 2025-06-01
---

# Python Best Practices

A comprehensive guide covering all aspects of Python development.
"""

    body_parts = [header]
    for title, content in sections:
        body_parts.append(f"\n## {title}\n\n{content}\n")

    return "\n".join(body_parts)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def repo_all_formats(tmp_path):
    """Build a temporary repo with test documents in all supported formats.

    Directory layout:
        finance/reports/2025-01-15-q4-revenue.md
        finance/invoices/2025-02-invoice.pdf
        finance/data/budget-2025.xlsx
        health/exercise-routine.md
        technical/guides/python-best-practices.md
        technical/specs/api-specification.docx

    Returns the tmp_path (repo root).
    """
    # -- Markdown files -----------------------------------------------------
    md_revenue = tmp_path / "finance" / "reports" / "2025-01-15-q4-revenue.md"
    md_revenue.parent.mkdir(parents=True, exist_ok=True)
    md_revenue.write_text(_Q4_REVENUE_MD, encoding="utf-8")

    md_exercise = tmp_path / "health" / "exercise-routine.md"
    md_exercise.parent.mkdir(parents=True, exist_ok=True)
    md_exercise.write_text(_EXERCISE_ROUTINE_MD, encoding="utf-8")

    md_python = tmp_path / "technical" / "guides" / "python-best-practices.md"
    md_python.parent.mkdir(parents=True, exist_ok=True)
    md_python.write_text(_generate_python_best_practices_md(), encoding="utf-8")

    # -- PDF ----------------------------------------------------------------
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_path = tmp_path / "finance" / "invoices" / "2025-02-invoice.pdf"
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.setFont("Helvetica-Bold", 18)
    c.drawString(72, 720, "Invoice #12345")
    c.setFont("Helvetica", 12)
    c.drawString(72, 690, "Date: 2025-02-01")
    c.drawString(72, 670, "Client: Acme Corporation")
    c.drawString(72, 650, "Amount: $5,000")
    c.drawString(72, 620, "Description: Consulting services for Q1 2025")
    c.drawString(72, 600, "Payment Terms: Net 30")
    c.save()

    # -- DOCX ---------------------------------------------------------------
    from docx import Document

    docx_path = tmp_path / "technical" / "specs" / "api-specification.docx"
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("API Specification", level=1)
    doc.add_paragraph(
        "This document describes the REST API for the platform. "
        "All endpoints require authentication via Bearer tokens."
    )
    doc.add_heading("Authentication", level=2)
    doc.add_paragraph(
        "Authentication uses Bearer tokens issued by the /auth/token endpoint. "
        "Include the token in the Authorization header of every request. "
        "Tokens expire after 3600 seconds and must be refreshed."
    )
    doc.add_heading("Endpoints", level=2)
    doc.add_paragraph(
        "The API exposes RESTful endpoints for managing resources. "
        "GET /api/v1/users returns a paginated list of users. "
        "POST /api/v1/users creates a new user account. "
        "PUT /api/v1/users/{id} updates an existing user. "
        "DELETE /api/v1/users/{id} deactivates a user account."
    )
    doc.save(str(docx_path))

    # -- XLSX ---------------------------------------------------------------
    from openpyxl import Workbook

    xlsx_path = tmp_path / "finance" / "data" / "budget-2025.xlsx"
    xlsx_path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Q1 Budget"
    ws1.append(["Category", "Amount", "Notes"])
    ws1.append(["Marketing", 50000, "Digital campaigns and events"])
    ws1.append(["Engineering", 120000, "Salaries and infrastructure"])
    ws1.append(["Operations", 30000, "Office and administration"])

    ws2 = wb.create_sheet("Q2 Budget")
    ws2.append(["Category", "Amount", "Notes"])
    ws2.append(["Marketing", 55000, "Expanded conference budget"])
    ws2.append(["Engineering", 125000, "New hire onboarding"])
    ws2.append(["Operations", 32000, "Office expansion costs"])

    wb.save(str(xlsx_path))

    return tmp_path


@pytest.fixture
def db_path(tmp_path):
    """Return a path for a fresh vector database."""
    return tmp_path / ".vectordb"


@pytest.fixture
def ingested_db(repo_all_formats, tmp_path_factory):
    """A fully ingested database from the all-formats repo."""
    from ingest import ingest

    db_path = tmp_path_factory.mktemp("vectordb")
    ingest(repo_root=repo_all_formats, db_path=db_path, force=True)
    return db_path
